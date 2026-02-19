#!/usr/bin/env python3
"""
End-to-End Test for Split-View Resume Editor

Tests the complete workflow:
1. Upload resume
2. Get sections detected
3. Update section content
4. Verify preview URL updated
5. Download updated DOCX
"""

import requests
import time
from pathlib import Path

# Configuration
BASE_URL = "http://localhost:8000"
UPLOAD_ENDPOINT = f"{BASE_URL}/api/upload"
PREVIEW_UPDATE_ENDPOINT = f"{BASE_URL}/api/preview/update"

# ANSI color codes
GREEN = '\033[92m'
RED = '\033[91m'
YELLOW = '\033[93m'
BLUE = '\033[94m'
RESET = '\033[0m'

def print_step(step_num, message):
    print(f"\n{BLUE}[Step {step_num}]{RESET} {message}")

def print_success(message):
    print(f"{GREEN}✓{RESET} {message}")

def print_error(message):
    print(f"{RED}✗{RESET} {message}")

def print_info(message):
    print(f"{YELLOW}ℹ{RESET} {message}")

def test_split_view_editor():
    """Run complete end-to-end test"""

    print(f"\n{'='*60}")
    print(f"{BLUE}Split-View Resume Editor - End-to-End Test{RESET}")
    print(f"{'='*60}")

    # Step 1: Create test resume file
    print_step(1, "Creating test DOCX resume...")

    from docx import Document
    test_resume = Document()
    test_resume.add_heading('John Doe', level=1)
    test_resume.add_paragraph('Software Engineer | Python Developer')
    test_resume.add_heading('Experience', level=2)
    test_resume.add_paragraph('Senior Developer at Tech Corp')
    test_resume.add_paragraph('Developed scalable web applications using Python and Django.')
    test_resume.add_heading('Education', level=2)
    test_resume.add_paragraph('BS Computer Science - MIT')
    test_resume.add_heading('Skills', level=2)
    test_resume.add_paragraph('Python, JavaScript, React, Django, PostgreSQL')

    test_file_path = Path('/tmp/test_resume_e2e.docx')
    test_resume.save(str(test_file_path))
    print_success(f"Created test resume at {test_file_path}")

    # Step 2: Upload resume
    print_step(2, "Uploading resume to backend...")

    with open(test_file_path, 'rb') as f:
        files = {'file': ('test_resume.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {
            'role': 'software-engineer',
            'level': 'mid',
            'jobDescription': 'Looking for a mid-level Python developer',
            'mode': 'ats'
        }

        try:
            response = requests.post(UPLOAD_ENDPOINT, files=files, data=data, timeout=30)

            if response.status_code != 200:
                print_error(f"Upload failed with status {response.status_code}")
                print_info(f"Response: {response.text[:500]}")
                return False

            upload_result = response.json()
            print_success(f"Upload successful (status: {response.status_code})")
        except Exception as e:
            print_error(f"Upload failed: {e}")
            if hasattr(e, 'response') and e.response:
                print_info(f"Response: {e.response.text[:500]}")
            return False

    # Step 3: Verify sections detected
    print_step(3, "Verifying sections detected...")

    if 'sections' not in upload_result:
        print_error("No sections found in response")
        print_info(f"Response keys: {upload_result.keys()}")
        return False

    sections = upload_result['sections']
    print_success(f"Detected {len(sections)} sections:")
    for i, section in enumerate(sections, 1):
        print(f"   {i}. {section['title']} (paragraphs {section['start_para_idx']}-{section['end_para_idx']})")

    # Step 4: Verify session ID and preview URL
    print_step(4, "Verifying session and preview URL...")

    if 'sessionId' not in upload_result:
        print_error("No sessionId in response")
        return False

    session_id = upload_result['sessionId']
    print_success(f"Session ID: {session_id}")

    if 'previewUrl' not in upload_result:
        print_error("No previewUrl in response")
        return False

    preview_url = upload_result['previewUrl']
    print_success(f"Preview URL: {preview_url}")

    # Step 5: Test section update
    print_step(5, "Testing section content update...")

    if len(sections) < 1:
        print_error("No sections to update")
        return False

    # Update the first section (Experience)
    first_section = sections[0]
    original_content = first_section['content']
    new_content = original_content + "\nAdditionally, led a team of 5 developers on microservices architecture."

    update_payload = {
        'session_id': session_id,
        'start_para_idx': first_section['start_para_idx'],
        'end_para_idx': first_section['end_para_idx'],
        'new_content': new_content
    }

    try:
        response = requests.post(PREVIEW_UPDATE_ENDPOINT, json=update_payload, timeout=10)
        response.raise_for_status()
        update_result = response.json()
        print_success(f"Section update successful (status: {response.status_code})")
    except Exception as e:
        print_error(f"Section update failed: {e}")
        if hasattr(e, 'response'):
            print_info(f"Response: {e.response.text}")
        return False

    # Step 6: Verify preview URL updated
    print_step(6, "Verifying preview URL updated...")

    if 'preview_url' not in update_result:
        print_error("No preview_url in update response")
        print_info(f"Response keys: {update_result.keys()}")
        return False

    new_preview_url = update_result['preview_url']

    if new_preview_url == preview_url:
        print_error("Preview URL did not change (cache-busting failed)")
        return False

    print_success(f"Preview URL updated with cache-busting: {new_preview_url}")

    # Step 7: Verify updated DOCX is accessible
    print_step(7, "Verifying updated DOCX is accessible...")

    try:
        response = requests.get(f"{BASE_URL}{new_preview_url}", timeout=10)
        response.raise_for_status()

        if response.headers.get('Content-Type') != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            print_error(f"Wrong content type: {response.headers.get('Content-Type')}")
            return False

        docx_size = len(response.content)
        print_success(f"DOCX accessible ({docx_size} bytes)")

        # Save and verify content
        updated_docx_path = Path('/tmp/test_resume_updated.docx')
        updated_docx_path.write_bytes(response.content)

        # Verify the update is in the DOCX
        from docx import Document
        updated_doc = Document(str(updated_docx_path))
        full_text = '\n'.join([p.text for p in updated_doc.paragraphs])

        if "led a team of 5 developers" in full_text:
            print_success("Updated content verified in DOCX file")
        else:
            print_error("Updated content not found in DOCX")
            print_info(f"Expected: 'led a team of 5 developers'")
            print_info(f"Full text preview: {full_text[:200]}...")
            return False

    except Exception as e:
        print_error(f"Failed to access updated DOCX: {e}")
        return False

    # Step 8: Test security validation (invalid session ID)
    print_step(8, "Testing security validation (invalid session ID)...")

    invalid_payload = {
        'session_id': '../../../etc/passwd',  # Path traversal attempt
        'start_para_idx': 0,
        'end_para_idx': 1,
        'new_content': 'malicious content'
    }

    try:
        response = requests.post(PREVIEW_UPDATE_ENDPOINT, json=invalid_payload, timeout=10)

        if response.status_code == 400:
            print_success("Security validation working (rejected invalid session ID)")
        else:
            print_error(f"Security issue: accepted invalid session ID (status: {response.status_code})")
            return False

    except Exception as e:
        print_error(f"Security test failed with exception: {e}")
        return False

    # Step 9: Test boundary validation (negative indices)
    print_step(9, "Testing boundary validation (negative indices)...")

    boundary_payload = {
        'session_id': session_id,
        'start_para_idx': -1,  # Negative index
        'end_para_idx': 5,
        'new_content': 'test content'
    }

    try:
        response = requests.post(PREVIEW_UPDATE_ENDPOINT, json=boundary_payload, timeout=10)
        result = response.json()

        if not result.get('success', True):
            print_success("Boundary validation working (rejected negative index)")
        else:
            print_error("Boundary issue: accepted negative index")
            return False

    except Exception as e:
        print_error(f"Boundary test failed with exception: {e}")
        return False

    # Summary
    print(f"\n{'='*60}")
    print(f"{GREEN}✓ All End-to-End Tests PASSED{RESET}")
    print(f"{'='*60}\n")

    print(f"{BLUE}Summary:{RESET}")
    print(f"  • Upload: {GREEN}✓{RESET} Working")
    print(f"  • Section Detection: {GREEN}✓{RESET} Working ({len(sections)} sections)")
    print(f"  • Section Update: {GREEN}✓{RESET} Working")
    print(f"  • Preview Generation: {GREEN}✓{RESET} Working (cache-busting enabled)")
    print(f"  • DOCX Download: {GREEN}✓{RESET} Working")
    print(f"  • Content Verification: {GREEN}✓{RESET} Working")
    print(f"  • Security Validation: {GREEN}✓{RESET} Working (UUID validation)")
    print(f"  • Boundary Validation: {GREEN}✓{RESET} Working (negative index check)")

    print(f"\n{GREEN}The split-view editor is fully functional and ready for production!{RESET}\n")

    return True


if __name__ == '__main__':
    try:
        success = test_split_view_editor()
        exit(0 if success else 1)
    except KeyboardInterrupt:
        print(f"\n{YELLOW}Test interrupted by user{RESET}")
        exit(130)
    except Exception as e:
        print(f"\n{RED}Test failed with unexpected error: {e}{RESET}")
        import traceback
        traceback.print_exc()
        exit(1)
