"""Manual test runner for template manager"""
import sys
sys.path.insert(0, '/Users/sabuj.mondal/ats-resume-scorer/backend')

from docx import Document
from io import BytesIO
import os
import tempfile
from services.docx_template_manager import DocxTemplateManager

def create_test_docx():
    """Create a test DOCX file"""
    doc = Document()
    doc.add_heading('Contact Information', level=2)
    doc.add_paragraph('John Doe')
    doc.add_paragraph('john@example.com')
    doc.add_heading('Experience', level=2)
    doc.add_paragraph('Software Engineer at ABC Corp')

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.read()

def test_save_template():
    """Test saving original template"""
    print("Running test_save_template...")

    # Create temp directory
    with tempfile.TemporaryDirectory() as tmp_path:
        template_manager = DocxTemplateManager(storage_dir=tmp_path)
        test_docx = create_test_docx()
        session_id = "test_session_123"

        # Save template
        template_path = template_manager.save_template(session_id, test_docx)

        # Verify file exists
        assert os.path.exists(template_path), f"Template path does not exist: {template_path}"
        assert session_id in template_path, f"Session ID not in path: {template_path}"
        assert template_path.endswith('_original.docx'), f"Path doesn't end with _original.docx: {template_path}"

        # Verify working copy created
        working_path = template_path.replace('_original.docx', '_working.docx')
        assert os.path.exists(working_path), f"Working copy does not exist: {working_path}"

        print("✓ test_save_template PASSED")
        return True

if __name__ == "__main__":
    try:
        test_save_template()
        print("\n✓ All tests passed!")
        sys.exit(0)
    except AssertionError as e:
        print(f"\n✗ Test failed: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"\n✗ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
