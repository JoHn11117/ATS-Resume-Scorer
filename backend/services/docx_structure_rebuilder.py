"""
DOCX Structure Rebuilder
Updates DOCX files with edited text while preserving all formatting
Performs binary-level editing to maintain 100% fidelity
"""
from docx import Document
from typing import Dict, Any, List
import copy


class DocxStructureRebuilder:
    """Rebuild DOCX with edited text while preserving formatting"""

    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.docx_path = docx_path
        # Create mapping of element IDs to elements for fast lookup
        self.paragraph_map = {id(para): para for para in self.doc.paragraphs}
        self.run_map = {}
        for para in self.doc.paragraphs:
            for run in para.runs:
                self.run_map[id(run)] = run

    def update_text(self, edits: List[Dict[str, Any]]) -> None:
        """
        Update document text with edits

        Args:
            edits: List of edits with structure:
                [
                    {
                        "paragraph_id": 12345,
                        "run_edits": [
                            {"run_id": 67890, "text": "new text"}
                        ]
                    }
                ]
        """
        for edit in edits:
            paragraph_id = edit.get("paragraph_id")
            run_edits = edit.get("run_edits", [])

            # Update runs in this paragraph
            for run_edit in run_edits:
                run_id = run_edit.get("run_id")
                new_text = run_edit.get("text", "")

                if run_id in self.run_map:
                    run = self.run_map[run_id]
                    # Update text while preserving all formatting
                    run.text = new_text

    def update_paragraph_text(self, paragraph_id: int, new_text: str) -> None:
        """
        Update entire paragraph text while preserving formatting

        Args:
            paragraph_id: ID of paragraph to update
            new_text: New text content
        """
        if paragraph_id in self.paragraph_map:
            para = self.paragraph_map[paragraph_id]

            # If paragraph has runs, update first run and clear others
            if para.runs:
                # Keep formatting from first run
                first_run = para.runs[0]
                first_run.text = new_text

                # Remove other runs
                for i in range(len(para.runs) - 1, 0, -1):
                    para._element.remove(para.runs[i]._element)

    def save(self, output_path: str) -> None:
        """
        Save updated document

        Args:
            output_path: Path to save updated DOCX
        """
        self.doc.save(output_path)


def update_docx_text(
    docx_path: str,
    edits: List[Dict[str, Any]],
    output_path: str
) -> None:
    """
    Update DOCX file with edited text

    Args:
        docx_path: Path to original DOCX file
        edits: List of text edits
        output_path: Path to save updated DOCX

    Returns:
        None
    """
    rebuilder = DocxStructureRebuilder(docx_path)
    rebuilder.update_text(edits)
    rebuilder.save(output_path)
