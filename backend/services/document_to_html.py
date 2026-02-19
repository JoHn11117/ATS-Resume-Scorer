"""
Convert documents (PDF/DOCX) to rich HTML with formatting preserved.

This module converts resume files to editable HTML that preserves:
- Text formatting (bold, italic, underline)
- Document structure (headings, paragraphs, lists)
- Tables and layout
- Fonts and styling
"""
import io
import logging
from typing import Dict
import mammoth
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)


def _enhance_html_formatting(html: str) -> str:
    """
    Post-process HTML to improve formatting and structure.

    Args:
        html: Raw HTML from Mammoth

    Returns:
        Enhanced HTML with better formatting
    """
    import re

    # Replace multiple line breaks with proper paragraph spacing
    html = re.sub(r'<br\s*/?\s*>(\s*<br\s*/?\s*>)+', '</p><p>', html)

    # Ensure headings have proper spacing
    html = re.sub(r'</h([1-6])>\s*<p>', r'</h\1><p>', html)

    # Add line breaks before bullet points for better spacing
    html = re.sub(r'</ul>\s*<p>', '</ul><br><p>', html)
    html = re.sub(r'</ol>\s*<p>', '</ol><br><p>', html)

    # Clean up empty paragraphs
    html = re.sub(r'<p>\s*</p>', '<p>&nbsp;</p>', html)

    # Preserve non-breaking spaces
    html = html.replace('\u00a0', '&nbsp;')

    return html


def docx_to_html(docx_bytes: bytes) -> str:
    """
    Convert DOCX to rich HTML with formatting preserved using Mammoth.

    Args:
        docx_bytes: DOCX file content as bytes

    Returns:
        HTML string with formatting preserved
    """
    try:
        # Use mammoth for much better formatting preservation
        # Custom style mapping for better preservation
        style_map = """
        p[style-name='Heading 1'] => h1:fresh
        p[style-name='Heading 2'] => h2:fresh
        p[style-name='Heading 3'] => h3:fresh
        p[style-name='Title'] => h1:fresh
        p[style-name='Subtitle'] => h2:fresh
        """

        result = mammoth.convert_to_html(
            io.BytesIO(docx_bytes),
            style_map=style_map
        )
        html = result.value

        # Log any warnings
        if result.messages:
            for message in result.messages:
                logger.warning(f"Mammoth conversion: {message}")

        # Post-process HTML to improve formatting
        html = _enhance_html_formatting(html)

        logger.info(f"Mammoth converted DOCX to HTML ({len(html)} chars)")
        return html

    except Exception as e:
        logger.error(f"Mammoth conversion failed: {e}, falling back to basic converter")
        # Fallback to basic converter if mammoth fails
        return _docx_to_html_fallback(docx_bytes)


def _docx_to_html_fallback(docx_bytes: bytes) -> str:
    """
    Fallback DOCX to HTML converter (basic).

    Args:
        docx_bytes: DOCX file content as bytes

    Returns:
        HTML string with basic formatting
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(io.BytesIO(docx_bytes))
    html_parts = []

    # Remove inline styles - let the editor's CSS handle it
    # Just output clean HTML structure

    # Track list context
    in_list = False
    current_list_type = None

    # Process paragraphs
    for para in doc.paragraphs:
        # Check if this is a list item
        is_list_item = False
        list_type = None

        # Check if paragraph has numbering
        if hasattr(para, '_p') and hasattr(para._p, 'pPr'):
            pPr = para._p.pPr
            if pPr is not None and hasattr(pPr, 'numPr'):
                numPr = pPr.numPr
                if numPr is not None:
                    is_list_item = True
                    # Determine if numbered or bulleted (simplified)
                    list_type = 'ol'  # Default to ordered
                    if para.style.name and 'bullet' in para.style.name.lower():
                        list_type = 'ul'

        # Close previous list if needed
        if not is_list_item and in_list:
            html_parts.append(f"</{current_list_type}>")
            in_list = False
            current_list_type = None

        # Open new list if needed
        if is_list_item and not in_list:
            in_list = True
            current_list_type = list_type
            html_parts.append(f"<{current_list_type}>")

        # Handle empty paragraphs
        if not para.text.strip():
            if not in_list:
                html_parts.append("<br>")
            continue

        # Determine tag
        if is_list_item:
            tag = "li"
        elif para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading ', '').strip()
            if level.isdigit():
                tag = f"h{min(int(level), 6)}"
            else:
                tag = "h2"
        else:
            tag = "p"

        # Get alignment style
        align_style = ""
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_style = ' style="text-align: center;"'
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_style = ' style="text-align: right;"'

        # Build paragraph with inline formatting
        text_html = ""
        for run in para.runs:
            text = run.text
            if not text:
                continue

            # Escape HTML entities
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # Apply formatting
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"

            # Apply font size if significantly different
            if run.font.size:
                size_pt = run.font.size.pt
                if size_pt > 16:
                    text = f'<span style="font-size: {size_pt}pt;">{text}</span>'
                elif size_pt < 10:
                    text = f'<span style="font-size: {size_pt}pt;">{text}</span>'

            # Apply font color if set
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                color = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
                text = f'<span style="color: {color};">{text}</span>'

            text_html += text

        html_parts.append(f"<{tag}{align_style}>{text_html}</{tag}>")

    # Close any open list
    if in_list:
        html_parts.append(f"</{current_list_type}>")

    # Process tables
    for table in doc.tables:
        html_parts.append("<table>")
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                cell_html = ""
                for para in cell.paragraphs:
                    if para.text.strip():
                        # Build cell content with formatting
                        text_html = ""
                        for run in para.runs:
                            text = run.text
                            if run.bold:
                                text = f"<strong>{text}</strong>"
                            if run.italic:
                                text = f"<em>{text}</em>"
                            text_html += text
                        cell_html += text_html + "<br>"
                html_parts.append(f"<td>{cell_html}</td>")
            html_parts.append("</tr>")
        html_parts.append("</table>")

    return "".join(html_parts)


def pdf_to_html(pdf_bytes: bytes) -> str:
    """
    Convert PDF to rich HTML (simplified - extracts text with basic formatting).

    Args:
        pdf_bytes: PDF file content as bytes

    Returns:
        HTML string with basic formatting
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    html_parts = []

    # Add basic styling
    html_parts.append("""
    <style>
        body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; }
        h1 { font-size: 24px; font-weight: bold; margin: 20px 0 10px 0; }
        h2 { font-size: 18px; font-weight: bold; margin: 15px 0 8px 0; }
        p { margin: 8px 0; }
    </style>
    """)

    for page in doc:
        # Extract text blocks with basic formatting detection
        blocks = page.get_text("dict")["blocks"]

        for block in blocks:
            if block.get("type") == 0:  # Text block
                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        size = span.get("size", 12)
                        flags = span.get("flags", 0)

                        # Check for bold (flag 16) and italic (flag 2)
                        is_bold = flags & 16
                        is_italic = flags & 2

                        if is_bold:
                            text = f"<strong>{text}</strong>"
                        if is_italic:
                            text = f"<em>{text}</em>"

                        line_text += text

                    if line_text.strip():
                        # Detect headings by font size
                        avg_size = sum(s.get("size", 12) for s in line.get("spans", [])) / max(len(line.get("spans", [])), 1)
                        if avg_size > 16:
                            html_parts.append(f"<h1>{line_text}</h1>")
                        elif avg_size > 14:
                            html_parts.append(f"<h2>{line_text}</h2>")
                        else:
                            html_parts.append(f"<p>{line_text}</p>")

    doc.close()
    return "".join(html_parts)
