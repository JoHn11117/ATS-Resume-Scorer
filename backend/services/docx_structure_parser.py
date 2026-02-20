"""
DOCX Structure Parser
Parses DOCX files and extracts editable text content with formatting metadata
Preserves document structure for binary-level editing
"""
from docx import Document
from docx.shared import Pt, RGBColor
from typing import List, Dict, Any, Optional
import base64
from io import BytesIO


class DocxStructureParser:
    """Parse DOCX structure for binary-level editing"""

    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.docx_path = docx_path

    def parse(self) -> Dict[str, Any]:
        """
        Parse DOCX and return structured representation

        Returns:
            Dict with document structure and editable text
        """
        structure = {
            "sections": self._parse_sections(),
            "metadata": self._extract_metadata(),
        }
        return structure

    def _extract_metadata(self) -> Dict[str, Any]:
        """Extract document-level metadata"""
        return {
            "page_width": self.doc.sections[0].page_width.inches if self.doc.sections else 8.5,
            "page_height": self.doc.sections[0].page_height.inches if self.doc.sections else 11,
            "margin_left": self.doc.sections[0].left_margin.inches if self.doc.sections else 1,
            "margin_right": self.doc.sections[0].right_margin.inches if self.doc.sections else 1,
            "margin_top": self.doc.sections[0].top_margin.inches if self.doc.sections else 1,
            "margin_bottom": self.doc.sections[0].bottom_margin.inches if self.doc.sections else 1,
        }

    def _parse_sections(self) -> List[Dict[str, Any]]:
        """Parse document sections"""
        sections = []
        current_section = {
            "paragraphs": [],
            "tables": []
        }

        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # Paragraph
                para = self._find_paragraph_by_element(element)
                if para:
                    current_section["paragraphs"].append(self._parse_paragraph(para))
            elif element.tag.endswith('tbl'):
                # Table
                table = self._find_table_by_element(element)
                if table:
                    current_section["tables"].append(self._parse_table(table))

        sections.append(current_section)
        return sections

    def _find_paragraph_by_element(self, element) -> Optional[Any]:
        """Find paragraph object by element"""
        for para in self.doc.paragraphs:
            if para._element == element:
                return para
        return None

    def _find_table_by_element(self, element) -> Optional[Any]:
        """Find table object by element"""
        for table in self.doc.tables:
            if table._element == element:
                return table
        return None

    def _parse_paragraph(self, para) -> Dict[str, Any]:
        """Parse paragraph with runs"""
        return {
            "id": id(para),
            "text": para.text,
            "style": para.style.name if para.style else None,
            "alignment": str(para.alignment) if para.alignment else None,
            "runs": [self._parse_run(run) for run in para.runs],
            "formatting": {
                "space_before": para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
                "space_after": para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0,
                "line_spacing": para.paragraph_format.line_spacing if para.paragraph_format.line_spacing else 1,
                "left_indent": para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 0,
                "right_indent": para.paragraph_format.right_indent.pt if para.paragraph_format.right_indent else 0,
            }
        }

    def _parse_run(self, run) -> Dict[str, Any]:
        """Parse run (text fragment with consistent formatting)"""
        return {
            "id": id(run),
            "text": run.text,
            "formatting": {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name if run.font.name else "Calibri",
                "font_size": run.font.size.pt if run.font.size else 11,
                "font_color": self._get_color(run.font.color) if run.font.color else "#000000",
                "highlight": str(run.font.highlight_color) if run.font.highlight_color else None,
            }
        }

    def _get_color(self, color) -> str:
        """Convert color to hex string"""
        if color.rgb:
            return f"#{color.rgb}"
        return "#000000"

    def _parse_table(self, table) -> Dict[str, Any]:
        """Parse table structure"""
        return {
            "id": id(table),
            "rows": len(table.rows),
            "cols": len(table.columns),
            "cells": [
                [self._parse_cell(cell) for cell in row.cells]
                for row in table.rows
            ]
        }

    def _parse_cell(self, cell) -> Dict[str, Any]:
        """Parse table cell"""
        return {
            "text": cell.text,
            "paragraphs": [self._parse_paragraph(para) for para in cell.paragraphs]
        }


def parse_docx_structure(docx_path: str) -> Dict[str, Any]:
    """
    Parse DOCX file and return structured representation

    Args:
        docx_path: Path to DOCX file

    Returns:
        Dict with document structure
    """
    parser = DocxStructureParser(docx_path)
    return parser.parse()
