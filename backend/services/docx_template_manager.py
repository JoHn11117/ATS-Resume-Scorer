"""
Manage DOCX templates for resume editing.
Stores original and working copies, handles section updates.
"""
from pathlib import Path
from docx import Document
import shutil
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

class DocxTemplateManager:
    """Manage DOCX templates with section-level editing"""

    def __init__(self, storage_dir: str = None):
        if storage_dir is None:
            storage_dir = Path(__file__).parent.parent / "storage" / "templates"
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)

    def save_template(self, session_id: str, docx_bytes: bytes) -> str:
        """
        Save original template and create working copy.

        Args:
            session_id: Unique session identifier
            docx_bytes: Original DOCX file bytes

        Returns:
            Path to original template file
        """
        # Save original
        original_path = self.storage_dir / f"{session_id}_original.docx"
        with open(original_path, 'wb') as f:
            f.write(docx_bytes)

        logger.info(f"Saved original template: {original_path}")

        # Create working copy
        working_path = self.storage_dir / f"{session_id}_working.docx"
        shutil.copy(original_path, working_path)

        logger.info(f"Created working copy: {working_path}")

        return str(original_path)

    def get_working_path(self, session_id: str) -> Path:
        """Get path to working DOCX"""
        return self.storage_dir / f"{session_id}_working.docx"

    def working_exists(self, session_id: str) -> bool:
        """Check if working copy exists"""
        return self.get_working_path(session_id).exists()

    def update_section(
        self,
        session_id: str,
        start_para_idx: int,
        end_para_idx: int,
        new_content: str
    ) -> dict:
        """
        Update section content while preserving formatting.

        Args:
            session_id: Session identifier
            start_para_idx: Start paragraph index
            end_para_idx: End paragraph index (inclusive)
            new_content: New text content (can include newlines)

        Returns:
            Dict with success status and preview URL
        """
        working_path = self.get_working_path(session_id)

        if not working_path.exists():
            return {'success': False, 'error': 'Session not found'}

        try:
            doc = Document(working_path)
            paragraphs = doc.paragraphs

            # Validate indices - check order and bounds
            if start_para_idx < 0 or end_para_idx < 0:
                return {'success': False, 'error': 'Paragraph indices cannot be negative'}

            if start_para_idx > end_para_idx:
                return {'success': False, 'error': 'Start index must be less than or equal to end index'}

            if start_para_idx >= len(paragraphs) or end_para_idx >= len(paragraphs):
                return {'success': False, 'error': 'Invalid paragraph indices'}

            # Get formatting from first paragraph in range
            first_para = paragraphs[start_para_idx]
            style = first_para.style

            # Delete old paragraphs (keep first to preserve formatting)
            for i in range(end_para_idx, start_para_idx, -1):
                p = paragraphs[i]._element
                p.getparent().remove(p)

            # Update first paragraph with new content
            first_para.text = ''

            # Split new content by newlines and add as separate paragraphs
            lines = new_content.split('\n')
            first_para.add_run(lines[0])

            # Add remaining lines as new paragraphs after first
            for line in lines[1:]:
                new_p = first_para.insert_paragraph_before(line)
                new_p.style = style

            # Save updated document
            doc.save(working_path)

            logger.info(f"Updated section in {session_id}: paragraphs {start_para_idx}-{end_para_idx}")

            # Generate preview URL with timestamp
            timestamp = int(datetime.now().timestamp())
            preview_url = f"/api/preview/{session_id}.docx?v={timestamp}"

            return {
                'success': True,
                'preview_url': preview_url
            }

        except Exception as e:
            logger.error(f"Failed to update section: {e}")
            return {'success': False, 'error': str(e)}
