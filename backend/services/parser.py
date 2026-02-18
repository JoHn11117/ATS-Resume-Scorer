"""
PDF and DOCX parser service for extracting structured data from resumes.
"""
import re
import io
import logging
from io import BytesIO
from typing import Dict, List, Optional
from pydantic import BaseModel, Field
import fitz  # PyMuPDF
import pypdf
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


# Pydantic models for structured resume data
class ContactInfo(BaseModel):
    """Contact information extracted from resume"""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    location: Optional[str] = None


class Experience(BaseModel):
    """Work experience entry"""
    title: str
    company: str
    location: Optional[str] = None
    startDate: Optional[str] = None
    endDate: Optional[str] = None
    description: Optional[str] = None
    achievements: List[str] = Field(default_factory=list)


class Education(BaseModel):
    """Education entry"""
    degree: str
    institution: str
    location: Optional[str] = None
    graduationDate: Optional[str] = None
    gpa: Optional[str] = None
    relevantCourses: List[str] = Field(default_factory=list)


class Certification(BaseModel):
    """Certification entry"""
    name: str
    issuer: Optional[str] = None
    date: Optional[str] = None
    expirationDate: Optional[str] = None


class ResumeMetadata(BaseModel):
    """Metadata about the parsed resume"""
    pageCount: int
    wordCount: int
    hasPhoto: bool = False
    fileFormat: str  # "pdf" or "docx"


class ResumeData(BaseModel):
    """Complete structured resume data"""
    fileName: str
    contact: Dict[str, Optional[str]]
    experience: List[Dict] = Field(default_factory=list)
    education: List[Dict] = Field(default_factory=list)
    skills: List[str] = Field(default_factory=list)
    certifications: List[Dict] = Field(default_factory=list)
    metadata: Dict


# Helper functions for text extraction
def extract_email(text: str) -> Optional[str]:
    """Extract email address from text using regex"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    match = re.search(email_pattern, text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number from text using regex"""
    # Matches formats: (123) 456-7890, 123-456-7890, 123.456.7890, +1 123 456 7890
    phone_patterns = [
        r'\+?1?\s*\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',
        r'\+?\d{1,3}[\s.-]?\d{3}[\s.-]?\d{3}[\s.-]?\d{4}',
    ]
    for pattern in phone_patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0)
    return None


def extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn URL or reference from text"""
    # Try full URL first
    linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
    match = re.search(linkedin_pattern, text, re.IGNORECASE)
    if match:
        return match.group(0)

    # Try "LinkedIn: username" or "LinkedIn Profile" patterns
    linkedin_ref = r'linkedin\s*(?:profile)?:?\s*([\w-]+)'
    match = re.search(linkedin_ref, text, re.IGNORECASE)
    if match:
        username = match.group(1) if match.group(1) not in ['profile', 'Profile'] else None
        if username:
            return f"linkedin.com/in/{username}"

    # Check if "LinkedIn" is mentioned in header (likely has hyperlink in PDF)
    if re.search(r'\blinkedin\b', text[:500], re.IGNORECASE):
        return "LinkedIn (see resume)"

    return None


def extract_location(text: str) -> Optional[str]:
    """Extract location (City, State/Country) from text"""
    # Common patterns: "City, State", "City, Country", "City | State"
    # Look for: Capital Letter(s), comma/pipe, Capital Letter(s)
    location_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*[,|]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'
    match = re.search(location_pattern, text)
    if match:
        return f"{match.group(1)}, {match.group(2)}"

    # Try simpler pattern: City Name (single or two words)
    # Look for lines that end with country/state names
    common_places = ['India', 'USA', 'UK', 'Canada', 'California', 'New York', 'Texas', 'London', 'Mumbai', 'Delhi', 'Bangalore', 'Gurugram', 'Hyderabad']
    for place in common_places:
        if re.search(rf'\b{place}\b', text, re.IGNORECASE):
            # Find the full location string around this place
            pattern = rf'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*[,|•]\s*{place}'
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return f"{match.group(1)}, {place}"
            # Just return the place itself if no city found
            return place

    return None


def extract_name_from_header(text: str) -> Optional[str]:
    """
    Extract name from resume header (first few lines).
    Assumes name is on the first line or early in the document.
    """
    lines = text.strip().split('\n')
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        # Name is usually a short line with 2-4 words, no special characters
        if line and len(line.split()) <= 4 and not re.search(r'[@\d()]', line):
            # Skip lines that are likely section headers
            if line.lower() not in ['resume', 'cv', 'curriculum vitae', 'contact', 'profile']:
                return line
    return None


def parse_experience_entry(text: str) -> Dict:
    """
    Parse an experience entry into structured fields.

    Args:
        text: Raw experience text

    Returns:
        Dict with title, company, location, dates, description
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]

    entry = {
        'title': '',
        'company': '',
        'location': '',
        'startDate': '',
        'endDate': '',
        'description': text  # Fallback to full text
    }

    if not lines:
        return entry

    # Smart detection: check if first line looks like a company or job title
    # Company indicators: ALL CAPS, multiple capitals, common company words
    # Job title indicators: contains role keywords like Manager, Engineer, Developer, etc.

    title_keywords = ['manager', 'engineer', 'developer', 'analyst', 'designer', 'director',
                      'coordinator', 'specialist', 'consultant', 'lead', 'senior', 'junior',
                      'associate', 'intern', 'executive', 'officer', 'head', 'chief', 'architect']

    first_line = lines[0]
    second_line = lines[1] if len(lines) > 1 else ""

    # Check if first line looks like a company (mostly uppercase or has location info)
    first_is_company = (
        first_line.isupper() or  # All caps like "AIR INDIA"
        sum(1 for c in first_line if c.isupper()) > len(first_line) * 0.5 or  # More than 50% caps
        ',' in first_line  # Has location separator
    )

    # Check if second line looks like a job title
    second_is_title = any(keyword in second_line.lower() for keyword in title_keywords)

    # Determine format
    if first_is_company or second_is_title:
        # Format: Company + Location on line 1, Job Title on line 2
        # Parse company and location from first line
        if ',' in first_line:
            parts = first_line.split(',', 1)
            entry['company'] = parts[0].strip()
            entry['location'] = parts[1].strip() if len(parts) > 1 else ''
        else:
            entry['company'] = first_line

        # Job title on second line
        entry['title'] = second_line

        # Look for dates in third line
        if len(lines) > 2:
            third_line = lines[2]
            date_pattern = r'(\w+\s+\d{4}|\d{4})\s*[-–]\s*(\w+\s+\d{4}|\d{4}|Present|Current)'
            date_match = re.search(date_pattern, third_line, re.IGNORECASE)
            if date_match:
                entry['startDate'] = date_match.group(1)
                entry['endDate'] = date_match.group(2)

        # Description starts from line 3 or 4
        if len(lines) > 3:
            entry['description'] = '\n'.join(lines[3:])
        elif len(lines) > 2:
            entry['description'] = '\n'.join(lines[2:])
    else:
        # Format: Job Title on line 1, Company on line 2
        entry['title'] = first_line

        # Parse company and location from second line
        if ' - ' in second_line:
            parts = second_line.split(' - ')
            entry['company'] = parts[0].strip()
            if len(parts) > 1:
                entry['location'] = parts[1].strip()
        elif ',' in second_line:
            parts = second_line.split(',')
            entry['company'] = parts[0].strip()
            if len(parts) > 1:
                entry['location'] = parts[1].strip()
        else:
            entry['company'] = second_line

        # Look for dates in third line
        if len(lines) > 2:
            third_line = lines[2]
            date_pattern = r'(\w+\s+\d{4}|\d{4})\s*[-–]\s*(\w+\s+\d{4}|\d{4}|Present|Current)'
            date_match = re.search(date_pattern, third_line, re.IGNORECASE)
            if date_match:
                entry['startDate'] = date_match.group(1)
                entry['endDate'] = date_match.group(2)

        # Description starts from line 3
        if len(lines) > 3:
            entry['description'] = '\n'.join(lines[3:])
        elif len(lines) > 2:
            entry['description'] = '\n'.join(lines[2:])

    return entry


def parse_education_entry(text: str) -> Dict:
    """
    Parse an education entry into structured fields.

    Args:
        text: Raw education text

    Returns:
        Dict with degree, institution, location, graduationDate
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]

    entry = {
        'degree': '',
        'institution': '',
        'location': '',
        'graduationDate': '',
        'gpa': ''
    }

    if not lines:
        return entry

    # First line is usually the degree
    if lines:
        entry['degree'] = lines[0]

    # Second line often has institution - location
    if len(lines) > 1:
        second_line = lines[1]
        if ' - ' in second_line:
            parts = second_line.split(' - ')
            entry['institution'] = parts[0].strip()
            if len(parts) > 1:
                entry['location'] = parts[1].strip()
        elif ',' in second_line:
            parts = second_line.split(',')
            entry['institution'] = parts[0].strip()
            if len(parts) > 1:
                entry['location'] = parts[1].strip()
        else:
            entry['institution'] = second_line

    # Look for graduation date
    if len(lines) > 2:
        for line in lines[2:]:
            # Match "Graduated: 2015" or "2015" or "May 2015"
            grad_pattern = r'(?:Graduated:?\s*)?(\w+\s+\d{4}|\d{4})'
            grad_match = re.search(grad_pattern, line, re.IGNORECASE)
            if grad_match:
                entry['graduationDate'] = grad_match.group(1)

            # Look for GPA
            gpa_pattern = r'GPA:?\s*([\d.]+)'
            gpa_match = re.search(gpa_pattern, line, re.IGNORECASE)
            if gpa_match:
                entry['gpa'] = gpa_match.group(1)

    return entry


def split_education_entries(text: str) -> List[str]:
    """
    Split multiple education entries from a single text block.

    Detects new entries by:
    - Degree keywords (Bachelor, Master, PhD, etc.)
    - Blank lines separating entries

    Args:
        text: Raw education section text with potentially multiple entries

    Returns:
        List of individual education entry texts
    """
    # Degree keywords that indicate a new education entry
    degree_keywords = [
        r'\b(Bachelor|B\.?\s*[ASC]\.?|BA|BS|BBA|BEng|BTech)\b',
        r'\b(Master|M\.?\s*[ASC]\.?|MA|MS|MBA|MEng|MTech|MSc)\b',
        r'\b(Doctor|PhD|Ph\.?D\.?|Doctorate)\b',
        r'\b(Associate|A\.?\s*[AS]\.?|AA|AS)\b',
        r'\b(Diploma|Certificate|Certification)\b',
        r'\b(High School|Secondary School|XII|X)\b'
    ]

    lines = text.split('\n')
    entries = []
    current_entry = []

    for line in lines:
        line_stripped = line.strip()

        # Skip empty lines
        if not line_stripped:
            continue

        # Check if this line starts a new education entry
        is_new_entry = False
        for pattern in degree_keywords:
            if re.search(pattern, line_stripped, re.IGNORECASE):
                is_new_entry = True
                break

        if is_new_entry and current_entry:
            # Save current entry and start new one
            entries.append('\n'.join(current_entry))
            current_entry = [line_stripped]
        else:
            # Add to current entry
            current_entry.append(line_stripped)

    # Add last entry
    if current_entry:
        entries.append('\n'.join(current_entry))

    return entries


def extract_resume_sections(text: str) -> Dict[str, List]:
    """
    Extract structured sections from resume text.

    Args:
        text: Full resume text

    Returns:
        Dictionary with sections: experience, education, skills, certifications
        - experience: List[Dict] with structured fields (title, company, dates, etc.)
        - education: List[Dict] with structured fields (degree, institution, dates, etc.)
        - skills: List[str] (strings directly)
        - certifications: List[Dict] with name field
    """
    sections = {
        'experience': [],
        'education': [],
        'skills': [],
        'certifications': []
    }

    lines = text.split('\n')
    current_section = None
    current_content = []

    # Section headers to detect
    experience_headers = ['experience', 'work history', 'employment', 'professional experience', 'work experience']
    education_headers = ['education', 'academic background', 'qualifications', 'academic']
    skills_headers = ['skills', 'technical skills', 'core competencies', 'expertise', 'technologies']
    cert_headers = ['certifications', 'certificates', 'licenses', 'professional certifications']

    # Debug: Log all section headers found
    logger.info(f"Scanning {len(lines)} lines for resume sections")

    for line in lines:
        line_lower = line.lower().strip()

        # Detect section headers
        if any(header in line_lower for header in experience_headers):
            if current_section and current_content:
                content_text = '\n'.join(current_content)
                if current_section == 'experience':
                    sections[current_section].append(parse_experience_entry(content_text))
                elif current_section == 'education':
                    sections[current_section].append(parse_education_entry(content_text))
                elif current_section == 'certifications':
                    sections[current_section].append({'name': content_text})
            current_section = 'experience'
            current_content = []
        elif any(header in line_lower for header in education_headers):
            if current_section and current_content:
                content_text = '\n'.join(current_content)
                if current_section == 'experience':
                    sections[current_section].append(parse_experience_entry(content_text))
                elif current_section == 'education':
                    sections[current_section].append(parse_education_entry(content_text))
                elif current_section == 'certifications':
                    sections[current_section].append({'name': content_text})
            current_section = 'education'
            current_content = []
        elif any(header in line_lower for header in skills_headers):
            logger.info(f"Found skills header: '{line}' (matched: {[h for h in skills_headers if h in line_lower]})")
            if current_section and current_content:
                content_text = '\n'.join(current_content)
                if current_section == 'experience':
                    sections[current_section].append(parse_experience_entry(content_text))
                elif current_section == 'education':
                    sections[current_section].append(parse_education_entry(content_text))
                elif current_section == 'certifications':
                    sections[current_section].append({'name': content_text})
            current_section = 'skills'
            current_content = []
        elif any(header in line_lower for header in cert_headers):
            if current_section and current_content:
                content_text = '\n'.join(current_content)
                if current_section == 'experience':
                    sections[current_section].append(parse_experience_entry(content_text))
                elif current_section == 'education':
                    sections[current_section].append(parse_education_entry(content_text))
                elif current_section == 'certifications':
                    sections[current_section].append({'name': content_text})
            current_section = 'certifications'
            current_content = []
        elif current_section and line.strip():
            current_content.append(line.strip())

    # Add last section
    if current_section and current_content:
        content_text = '\n'.join(current_content)
        if current_section == 'experience':
            sections[current_section].append(parse_experience_entry(content_text))
        elif current_section == 'education':
            # Split multiple education entries by blank lines or degree patterns
            edu_entries = split_education_entries(content_text)
            for edu_text in edu_entries:
                if edu_text.strip():
                    sections[current_section].append(parse_education_entry(edu_text))
        elif current_section == 'skills':
            # For skills, just add the raw text to be processed below
            logger.info(f"Adding skills content: {content_text[:200]}...")
            sections[current_section].append(content_text)
        elif current_section == 'certifications':
            sections[current_section].append({'name': content_text})

    # Special handling for skills - split by commas/bullets (keep as strings)
    logger.info(f"Skills section raw data: {sections['skills']}")
    if sections['skills']:
        all_skills = []
        for skill_block in sections['skills']:
            # Split by common delimiters
            skills = re.split(r'[,;•|\n]+', skill_block)
            all_skills.extend([s.strip() for s in skills if s.strip()])
        sections['skills'] = list(set(all_skills))[:50]  # Deduplicate, limit to 50
        logger.info(f"Processed {len(sections['skills'])} skills: {sections['skills'][:10]}")
    else:
        logger.warning("No skills section found in resume")

    return sections


def parse_pdf_with_pypdf(file_content: bytes, filename: str) -> ResumeData:
    """
    Parse PDF using pypdf library (fallback strategy).

    Args:
        file_content: PDF file bytes
        filename: Original filename

    Returns:
        ResumeData with extracted content
    """
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_content))

        # Extract text from all pages
        full_text = ""
        for page in reader.pages:
            full_text += page.extract_text() + "\n"

        # Get metadata
        page_count = len(reader.pages)
        word_count = len(full_text.split())

        # Extract sections
        sections = extract_resume_sections(full_text)

        # Extract contact info
        header_text = full_text[:500]
        name = extract_name_from_header(full_text)
        email = extract_email(header_text)
        phone = extract_phone(header_text)
        linkedin = extract_linkedin(header_text)

        contact_info = {
            "name": name,
            "email": email,
            "phone": phone,
            "linkedin": linkedin
        }

        metadata = {
            "pageCount": page_count,
            "wordCount": word_count,
            "hasPhoto": False,
            "fileFormat": "pdf"
        }

        return ResumeData(
            fileName=filename,
            contact=contact_info,
            experience=sections.get('experience', []),
            education=sections.get('education', []),
            skills=sections.get('skills', []),
            certifications=sections.get('certifications', []),
            metadata=metadata
        )

    except Exception as e:
        # Return minimal result on failure
        return ResumeData(
            fileName=filename,
            contact={},
            experience=[],
            education=[],
            skills=[],
            metadata={"pageCount": 0, "wordCount": 0, "hasPhoto": False, "fileFormat": "pdf"}
        )


def parse_pdf_with_pdfplumber(file_content: bytes, filename: str) -> ResumeData:
    """
    Parse PDF using pdfplumber library (fallback for tables).

    Args:
        file_content: PDF file bytes
        filename: Original filename

    Returns:
        ResumeData with extracted content
    """
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            # Extract text from all pages
            full_text_parts = []

            for page in pdf.pages:
                # Get regular text
                page_text = page.extract_text()
                if page_text:
                    full_text_parts.append(page_text)

                # Get table content (pdfplumber's specialty)
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            row_text = " | ".join([str(cell) for cell in row if cell])
                            full_text_parts.append(row_text)

            full_text = "\n".join(full_text_parts)

            # Get metadata
            page_count = len(pdf.pages)
            word_count = len(full_text.split())

            # Extract sections
            sections = extract_resume_sections(full_text)

            # Extract contact info
            header_text = full_text[:500]
            name = extract_name_from_header(full_text)
            email = extract_email(header_text)
            phone = extract_phone(header_text)
            linkedin = extract_linkedin(header_text)

            contact_info = {
                "name": name,
                "email": email,
                "phone": phone,
                "linkedin": linkedin,
                "location": extract_location(header_text)
            }

            metadata = {
                "pageCount": page_count,
                "wordCount": word_count,
                "hasPhoto": False,
                "fileFormat": "pdf"
            }

            return ResumeData(
                fileName=filename,
                contact=contact_info,
                experience=sections.get('experience', []),
                education=sections.get('education', []),
                skills=sections.get('skills', []),
                certifications=sections.get('certifications', []),
                metadata=metadata
            )

    except Exception as e:
        # Return minimal result on failure
        return ResumeData(
            fileName=filename,
            contact={},
            experience=[],
            education=[],
            skills=[],
            metadata={"pageCount": 0, "wordCount": 0, "hasPhoto": False, "fileFormat": "pdf"}
        )


def assess_parse_quality(resume: ResumeData, raw_text: str) -> float:
    """
    Assess quality of parsed resume.

    Args:
        resume: Parsed ResumeData
        raw_text: Raw extracted text

    Returns:
        Quality score 0.0-1.0
    """
    score = 0.0

    # Check word count
    if resume.metadata["wordCount"] >= 200:
        score += 0.3
    elif resume.metadata["wordCount"] >= 100:
        score += 0.15

    # Check sections found
    if resume.experience:
        score += 0.3
    if resume.education:
        score += 0.2
    if resume.skills:
        score += 0.2

    return score


def parse_pdf(file_content: bytes, filename: str) -> ResumeData:
    """
    Parse a PDF resume using multi-strategy approach.

    Tries PyMuPDF first, falls back to pypdf, then pdfplumber if needed.

    Args:
        file_content: PDF file content as bytes
        filename: Original filename of the PDF

    Returns:
        ResumeData object with extracted information
    """
    # Strategy 1: PyMuPDF (current implementation - fast and reliable)
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")

        # Extract text from all pages
        full_text = ""
        for page in doc:
            full_text += page.get_text()

        page_count = len(doc)
        word_count = len(full_text.split())
        doc.close()

        # If extraction seems successful, continue with PyMuPDF
        if word_count >= 50:  # Minimum threshold for valid extraction
            # Debug: log first 1000 chars of extracted text
            logger.info(f"PyMuPDF extracted {word_count} words from PDF")
            logger.info(f"First 1000 chars: {full_text[:1000]}")

            sections = extract_resume_sections(full_text)
            logger.info(f"Sections extracted - Experience: {len(sections.get('experience', []))}, Education: {len(sections.get('education', []))}, Skills: {len(sections.get('skills', []))}")
            logger.info(f"Education entries: {sections.get('education', [])}")
            logger.info(f"Skills entries: {sections.get('skills', [])}")

            # Debug: Check if skills section exists in text
            text_lower = full_text.lower()
            skills_found = 'skills' in text_lower or 'technical' in text_lower or 'competencies' in text_lower
            logger.info(f"Skills keywords in text: {skills_found}")

            header_text = full_text[:500]
            name = extract_name_from_header(full_text)
            email = extract_email(header_text)
            phone = extract_phone(header_text)
            linkedin = extract_linkedin(header_text)

            contact_info = {
                "name": name,
                "email": email,
                "phone": phone,
                "linkedin": linkedin,
                "location": extract_location(header_text)
            }

            metadata = {
                "pageCount": page_count,
                "wordCount": word_count,
                "hasPhoto": False,
                "fileFormat": "pdf"
            }

            result = ResumeData(
                fileName=filename,
                contact=contact_info,
                experience=sections.get('experience', []),
                education=sections.get('education', []),
                skills=sections.get('skills', []),
                certifications=sections.get('certifications', []),
                metadata=metadata
            )

            quality = assess_parse_quality(result, full_text)
            # If skills or education sections are missing, try pdfplumber (better for tables)
            if not sections.get('skills') or not sections.get('education'):
                logger.info(f"PyMuPDF missing skills/education sections, will try pdfplumber for table extraction")
            elif quality >= 0.7:  # Good quality, use it
                return result
    except Exception as e:
        pass  # Fall through to next strategy

    # Strategy 2: pypdf fallback
    try:
        result = parse_pdf_with_pypdf(file_content, filename)
        quality = assess_parse_quality(result, "")  # Note: no raw text available
        if quality >= 0.5:  # Lower threshold for fallback
            return result
    except Exception as e:
        pass  # Fall through to next strategy

    # Strategy 3: pdfplumber (best for tables)
    try:
        result = parse_pdf_with_pdfplumber(file_content, filename)
        return result  # Use whatever we got
    except Exception as e:
        # All strategies failed - return minimal result
        return ResumeData(
            fileName=filename,
            contact={},
            experience=[],
            education=[],
            skills=[],
            metadata={"pageCount": 0, "wordCount": 0, "hasPhoto": False, "fileFormat": "pdf"}
        )


def parse_docx(file_content: bytes, filename: str) -> ResumeData:
    """
    Parse a DOCX resume and extract structured data including tables.

    Args:
        file_content: DOCX file content as bytes
        filename: Original filename of the DOCX

    Returns:
        ResumeData object with extracted information
    """
    # Open DOCX from bytes using BytesIO
    doc = Document(BytesIO(file_content))

    # Extract text from paragraphs AND tables (CRITICAL FIX)
    full_text_parts = []

    # Get all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)

    # Get all tables (THIS WAS MISSING - CAUSED EMPTY SECTIONS)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text_parts.append(" | ".join(row_text))

    full_text = "\n".join(full_text_parts)

    # Extract sections (experience, education, skills)
    sections = extract_resume_sections(full_text)

    # Get metadata
    word_count = len(full_text.split())

    # Extract contact information from the header
    header_text = full_text[:500]

    name = extract_name_from_header(full_text)
    email = extract_email(header_text)
    phone = extract_phone(header_text)
    linkedin = extract_linkedin(header_text)

    contact_info = {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin
    }

    # Metadata
    metadata = {
        "pageCount": len(doc.sections),  # Approximate page count
        "wordCount": word_count,
        "hasPhoto": False,  # TODO: Implement image detection
        "fileFormat": "docx"
    }

    # Create ResumeData object with extracted sections
    resume_data = ResumeData(
        fileName=filename,
        contact=contact_info,
        experience=sections.get('experience', []),
        education=sections.get('education', []),
        skills=sections.get('skills', []),
        certifications=sections.get('certifications', []),
        metadata=metadata
    )

    return resume_data
