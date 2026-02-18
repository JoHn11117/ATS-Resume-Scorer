"""
DOCX to PDF converter for preview generation.

This module converts DOCX files to PDF format for consistent preview
in the editor. Uses python-docx to read DOCX and reportlab to generate PDF.
"""
import io
from typing import BinaryIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert DOCX file to PDF for preview.

    Args:
        docx_bytes: DOCX file content as bytes

    Returns:
        PDF file content as bytes
    """
    # Read DOCX
    doc = Document(io.BytesIO(docx_bytes))

    # Create PDF buffer
    pdf_buffer = io.BytesIO()
    pdf_doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )

    # Build story (content)
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='black',
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='black',
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        textColor='black',
        spaceAfter=6,
        fontName='Helvetica'
    )

    # Extract paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 0.1*inch))
            continue

        # Determine style based on text characteristics
        if i == 0 and len(text.split()) <= 4:
            # First short line is likely the name (title)
            story.append(Paragraph(text, title_style))
        elif text.isupper() or (len(text.split()) <= 4 and para.runs and para.runs[0].bold):
            # All caps or short bold text = heading
            story.append(Paragraph(text, heading_style))
        else:
            # Normal text
            story.append(Paragraph(text, normal_style))

    # Extract tables
    for table in doc.tables:
        story.append(Spacer(1, 0.2*inch))
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                story.append(Paragraph(row_text, normal_style))
        story.append(Spacer(1, 0.2*inch))

    # Build PDF
    pdf_doc.build(story)

    # Get PDF bytes
    pdf_buffer.seek(0)
    return pdf_buffer.read()
