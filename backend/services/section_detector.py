"""
Dynamically detect sections from DOCX resumes.
Supports various formats: heading styles, bold text, ALL CAPS, tables.
"""
from docx import Document
from docx.text.paragraph import Paragraph
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

class SectionDetector:
    """Detect resume sections dynamically without hardcoded section names"""

    # Constants for heading detection
    MIN_HEADING_FONT_SIZE = 12  # Minimum font size (in points) for bold text headings
    MIN_HEADING_LENGTH = 2      # Minimum length for ALL CAPS headings

    def detect(self, docx_bytes: bytes) -> list[dict]:
        """
        Detect sections from DOCX bytes.

        Args:
            docx_bytes: DOCX file content as bytes

        Returns:
            List of section dictionaries, each containing:
                - title (str): Section heading text
                - content (str): Section body content
                - section_id (str): Unique identifier (e.g., 'section_0')
                - start_para_idx (int): Starting paragraph index
                - end_para_idx (int): Ending paragraph index (exclusive)
                - is_in_table (bool): Whether section is in a table
                - table_cell_ref (str|None): Table cell reference if applicable

        Raises:
            ValueError: If docx_bytes is invalid or empty
        """
        if not docx_bytes:
            raise ValueError("docx_bytes cannot be empty")

        try:
            doc = Document(BytesIO(docx_bytes))
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise ValueError(f"Invalid DOCX format: {e}")

        sections = []
        current_section = None
        section_counter = 0

        # Handle empty document
        if not doc.paragraphs:
            logger.warning("Document has no paragraphs")
            return sections

        logger.info(f"Processing document with {len(doc.paragraphs)} paragraphs")

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a section heading
            is_heading = self._is_section_heading(para)

            if is_heading:
                # Save previous section
                if current_section:
                    current_section['end_para_idx'] = idx
                    sections.append(current_section)

                # Start new section
                current_section = {
                    'title': text,
                    'content': '',
                    'section_id': f'section_{section_counter}',
                    'start_para_idx': idx + 1,
                    'end_para_idx': idx + 1,
                    'is_in_table': False,
                    'table_cell_ref': None
                }
                section_counter += 1
            elif current_section:
                # Add content to current section
                if current_section['content']:
                    current_section['content'] += '\n'
                current_section['content'] += text
                current_section['end_para_idx'] = idx + 1

        # Save last section
        if current_section:
            sections.append(current_section)

        logger.info(f"Detected {len(sections)} sections")
        return sections

    def _is_section_heading(self, paragraph: Paragraph) -> bool:
        """
        Determine if paragraph is a section heading.

        Args:
            paragraph: DOCX paragraph object

        Returns:
            True if paragraph appears to be a section heading, False otherwise
        """
        # Check style name
        if 'Heading' in paragraph.style.name:
            logger.debug(f"Heading detected by style: {paragraph.text[:50]}")
            return True

        # Check if text is bold and larger font
        if paragraph.runs:
            first_run = paragraph.runs[0]
            try:
                # Safely check font size with proper error handling
                if first_run.bold and first_run.font.size and hasattr(first_run.font.size, 'pt'):
                    if first_run.font.size.pt >= self.MIN_HEADING_FONT_SIZE:
                        logger.debug(f"Heading detected by bold+size: {paragraph.text[:50]}")
                        return True
            except (AttributeError, TypeError) as e:
                # Font size may not be available for all runs
                logger.debug(f"Could not check font size: {e}")

        # Check for ALL CAPS (likely a heading)
        text = paragraph.text.strip()
        if text and text.isupper() and len(text) > self.MIN_HEADING_LENGTH:
            logger.debug(f"Heading detected by ALL CAPS: {text[:50]}")
            return True

        return False

    # Common section headers for keyword-based detection
    # Expanded to recognize many alternate names people use for sections
    SECTION_KEYWORDS = {
        'Contact': ['contact', 'personal', 'contact information', 'personal information', 'details'],
        'Summary': [
            'summary', 'objective', 'about', 'profile', 'brief',
            'professional summary', 'career summary', 'executive summary',
            'profile brief', 'professional profile', 'career profile',
            'about me', 'career objective', 'professional objective',
            'personal statement', 'introduction', 'overview',
            'professional overview', 'career overview'
        ],
        'Experience': [
            'experience', 'employment', 'work history', 'professional experience',
            'work experience', 'employment history', 'career history',
            'professional history', 'relevant experience', 'work'
        ],
        'Education': [
            'education', 'academic', 'qualifications', 'academic background',
            'educational background', 'academic qualifications', 'academics',
            'educational qualifications', 'training'
        ],
        'Skills': [
            'skills', 'technical skills', 'competencies', 'core competencies',
            'areas of expertise', 'key skills', 'expertise', 'capabilities',
            'technical competencies', 'professional skills', 'skillset'
        ],
        'Projects': [
            'projects', 'portfolio', 'key projects', 'notable projects',
            'project portfolio', 'major projects', 'project experience'
        ],
        'Certifications': [
            'certifications', 'certificates', 'licenses', 'professional certifications',
            'licensed', 'certification', 'professional licenses'
        ],
        'Awards': [
            'awards', 'honors', 'achievements', 'accomplishments',
            'recognition', 'honors and awards', 'awards and honors'
        ]
    }

    def detect_sections(self, doc: Document) -> list[dict]:
        """
        Detect sections in resume DOCX.

        Args:
            doc: python-docx Document object

        Returns:
            List of sections with name, start_para, end_para
        """
        sections = []
        paragraphs = doc.paragraphs

        # First section is usually contact (before first header)
        first_header_idx = None
        for i, para in enumerate(paragraphs):
            if self._is_section_header_keyword(para.text):
                first_header_idx = i
                break

        if first_header_idx is not None and first_header_idx > 0:
            sections.append({
                'name': 'Contact',
                'start_para': 0,
                'end_para': first_header_idx - 1
            })

        # Detect other sections by headers
        current_section = None
        start_idx = None

        for i, para in enumerate(paragraphs):
            section_name = self._identify_section(para.text)

            if section_name:
                # Save previous section
                if current_section and start_idx is not None:
                    sections.append({
                        'name': current_section,
                        'start_para': start_idx,
                        'end_para': i - 1
                    })

                # Start new section
                current_section = section_name
                start_idx = i

        # Save last section
        if current_section and start_idx is not None:
            sections.append({
                'name': current_section,
                'start_para': start_idx,
                'end_para': len(paragraphs) - 1
            })

        return sections

    def _is_section_header_keyword(self, text: str) -> bool:
        """Check if text looks like a section header"""
        if not text or len(text) > 50:
            return False

        text_lower = text.lower().strip()

        # Check against known section keywords
        for keywords in self.SECTION_KEYWORDS.values():
            for keyword in keywords:
                if keyword in text_lower:
                    return True

        return False

    def _identify_section(self, text: str) -> str:
        """Identify which section a header belongs to"""
        if not text:
            return None

        text_lower = text.lower().strip()

        for section_name, keywords in self.SECTION_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return section_name

        return None
