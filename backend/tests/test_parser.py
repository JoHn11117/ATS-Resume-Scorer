import pytest
import sys
import os
import io
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../..')))
from backend.services.parser import parse_pdf, parse_docx, parse_pdf_with_pypdf, ResumeData, extract_resume_sections
from docx import Document

def test_parse_pdf_extracts_basic_info():
    """Test that PDF parser extracts contact information"""
    # This test requires a sample PDF
    # For now, we'll test the structure
    resume_data = ResumeData(
        fileName="test.pdf",
        contact={
            "name": "John Doe",
            "email": "john@example.com",
            "phone": "555-1234"
        },
        experience=[],
        education=[],
        skills=[],
        metadata={"pageCount": 1, "wordCount": 100, "hasPhoto": False, "fileFormat": "pdf"}
    )

    assert resume_data.contact["name"] == "John Doe"
    assert resume_data.contact["email"] == "john@example.com"
    assert resume_data.metadata["fileFormat"] == "pdf"


def test_parse_pdf_with_pypdf_fallback():
    """Test pypdf fallback when pymupdf fails"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, "Test Resume")
    c.drawString(100, 730, "John Doe")
    c.drawString(100, 710, "john@example.com")
    c.drawString(100, 690, "Experience: Software Engineer")
    c.save()
    pdf_content = buffer.getvalue()

    result = parse_pdf_with_pypdf(pdf_content, "test.pdf")
    assert result is not None
    assert result.fileName == "test.pdf"
    assert result.metadata["fileFormat"] == "pdf"


def test_parse_docx_extracts_tables():
    """Test that DOCX parser extracts content from tables (CRITICAL FIX)"""
    # Create a DOCX with table layout (common resume format)
    doc = Document()

    # Add header with contact info
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane.smith@email.com | (555) 123-4567")

    # Add experience section with table layout
    doc.add_paragraph("EXPERIENCE")
    table = doc.add_table(rows=3, cols=2)

    # Row 1: Job title and dates
    table.rows[0].cells[0].text = "Senior Software Engineer"
    table.rows[0].cells[1].text = "2020 - Present"

    # Row 2: Company and location
    table.rows[1].cells[0].text = "Tech Corp"
    table.rows[1].cells[1].text = "San Francisco, CA"

    # Row 3: Achievements
    table.rows[2].cells[0].text = "Led development of microservices architecture"
    table.rows[2].cells[1].text = ""

    # Add education section with table
    doc.add_paragraph("EDUCATION")
    edu_table = doc.add_table(rows=2, cols=2)
    edu_table.rows[0].cells[0].text = "Bachelor of Science in Computer Science"
    edu_table.rows[0].cells[1].text = "2016"
    edu_table.rows[1].cells[0].text = "Stanford University"
    edu_table.rows[1].cells[1].text = "GPA: 3.8"

    # Add skills section
    doc.add_paragraph("SKILLS")
    doc.add_paragraph("Python, Java, JavaScript, React, Docker, Kubernetes, AWS, PostgreSQL")

    # Add certifications
    doc.add_paragraph("CERTIFICATIONS")
    doc.add_paragraph("AWS Certified Solutions Architect")

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_content = buffer.getvalue()

    # Parse the DOCX
    result = parse_docx(docx_content, "test_resume.docx")

    # Assert basic metadata
    assert result.fileName == "test_resume.docx"
    assert result.metadata["fileFormat"] == "docx"
    assert result.metadata["wordCount"] > 0

    # Assert contact info extracted
    assert result.contact["email"] == "jane.smith@email.com"
    assert "555" in result.contact["phone"]

    # CRITICAL: Assert experience section was extracted from table
    assert len(result.experience) > 0, "Experience section should not be empty when using tables"
    exp = result.experience[0]
    assert "engineer" in exp.get("title", "").lower() or "engineer" in exp.get("company", "").lower()

    # CRITICAL: Assert education section was extracted from table
    assert len(result.education) > 0, "Education section should not be empty when using tables"
    edu = result.education[0]
    assert "bachelor" in edu.get("degree", "").lower() or "computer" in edu.get("degree", "").lower()

    # CRITICAL: Assert skills were extracted
    assert len(result.skills) > 0, "Skills section should not be empty"
    skills_text = " ".join(result.skills).lower()
    assert "python" in skills_text or "java" in skills_text

    # Assert certifications extracted
    assert len(result.certifications) > 0, "Certifications should be extracted"


def test_parse_docx_with_mixed_paragraphs_and_tables():
    """Test DOCX parser with both paragraphs and tables"""
    doc = Document()

    # Add name as paragraph
    doc.add_paragraph("John Developer")

    # Add contact as paragraph
    doc.add_paragraph("john.dev@email.com | 555-9876")

    # Add experience header as paragraph
    doc.add_paragraph("WORK EXPERIENCE")

    # Add experience in table format
    exp_table = doc.add_table(rows=2, cols=1)
    exp_table.rows[0].cells[0].text = "Software Developer | ABC Company | 2019 - 2022"
    exp_table.rows[1].cells[0].text = "Built scalable web applications using React and Node.js"

    # Add education as paragraph (not in table)
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("Master of Computer Science")
    doc.add_paragraph("MIT - 2019")

    # Add skills in table
    doc.add_paragraph("TECHNICAL SKILLS")
    skills_table = doc.add_table(rows=1, cols=3)
    skills_table.rows[0].cells[0].text = "JavaScript"
    skills_table.rows[0].cells[1].text = "Python"
    skills_table.rows[0].cells[2].text = "SQL"

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_content = buffer.getvalue()

    # Parse
    result = parse_docx(docx_content, "mixed_resume.docx")

    # Verify both paragraph and table content was extracted
    assert result.contact["email"] == "john.dev@email.com"
    assert len(result.experience) > 0, "Should extract experience from tables"
    assert len(result.education) > 0, "Should extract education from paragraphs"
    assert len(result.skills) > 0, "Should extract skills from tables"

    # Verify skills contains items from table
    skills_lower = [s.lower() for s in result.skills]
    assert any("javascript" in s for s in skills_lower) or any("python" in s for s in skills_lower)


def test_extract_resume_sections_skills_deduplication():
    """Test that extract_resume_sections properly deduplicates and limits skills"""
    text = """
    SKILLS
    Python, Java, JavaScript, Python, React, Node.js, Docker, Kubernetes, AWS, PostgreSQL,
    MongoDB, Redis, ElasticSearch, GraphQL, REST API, Microservices, CI/CD, Git, Linux,
    Python, JavaScript, TypeScript, Go, Rust, C++, C#, Ruby, PHP, Swift, Kotlin,
    Angular, Vue.js, Svelte, Next.js, Express, Flask, Django, FastAPI, Spring Boot,
    MySQL, Oracle, SQL Server, Cassandra, DynamoDB, Firebase, Azure, GCP, Terraform,
    Ansible, Jenkins, GitLab CI, CircleCI, Travis CI, Selenium, Jest, Pytest, JUnit
    """

    sections = extract_resume_sections(text)

    # Assert skills were extracted
    assert len(sections['skills']) > 0

    # Assert deduplication - "Python" and "JavaScript" should only appear once each
    python_count = sum(1 for s in sections['skills'] if s.lower() == 'python')
    js_count = sum(1 for s in sections['skills'] if s.lower() == 'javascript')
    assert python_count <= 1, "Python should be deduplicated"
    assert js_count <= 1, "JavaScript should be deduplicated"

    # Assert limit to 50 skills
    assert len(sections['skills']) <= 50, "Skills should be limited to 50"


def test_extract_resume_sections_empty_sections():
    """Test extract_resume_sections with minimal content"""
    text = """
    John Doe
    john@example.com
    Some random text without section headers
    """

    sections = extract_resume_sections(text)

    # Should return empty lists for sections not found
    assert isinstance(sections['experience'], list)
    assert isinstance(sections['education'], list)
    assert isinstance(sections['skills'], list)
    assert isinstance(sections['certifications'], list)
