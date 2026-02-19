import pytest
from docx import Document
from io import BytesIO
import os
from pathlib import Path
from backend.services.docx_template_manager import DocxTemplateManager

@pytest.fixture
def test_docx():
    """Create a test DOCX file"""
    doc = Document()
    doc.add_heading('Contact Information', level=2)
    doc.add_paragraph('John Doe')
    doc.add_paragraph('john@example.com')
    doc.add_heading('Experience', level=2)
    doc.add_paragraph('Software Engineer at ABC Corp')

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.read()

@pytest.fixture
def template_manager(tmp_path):
    """Create template manager with temp storage"""
    return DocxTemplateManager(storage_dir=str(tmp_path))

def test_save_template(test_docx, template_manager):
    """Test saving original template"""
    session_id = "test_session_123"

    # Save template
    template_path = template_manager.save_template(session_id, test_docx)

    # Verify file exists
    assert os.path.exists(template_path)
    assert session_id in template_path
    assert template_path.endswith('_original.docx')

    # Verify working copy created
    working_path = template_path.replace('_original.docx', '_working.docx')
    assert os.path.exists(working_path)

def test_update_section_content(test_docx, template_manager):
    """Test updating specific section content"""
    session_id = "test_session_456"
    template_manager.save_template(session_id, test_docx)

    # Update Experience section (start_para_idx=4, end_para_idx=4)
    new_content = "Senior Software Engineer at XYZ Corp\nLed team of 5 developers"

    result = template_manager.update_section(
        session_id=session_id,
        start_para_idx=4,
        end_para_idx=4,
        new_content=new_content
    )

    assert result['success'] is True
    assert 'preview_url' in result

    # Verify content updated
    working_path = template_manager.get_working_path(session_id)
    doc = Document(working_path)

    # Check that new content is in document
    all_text = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Senior Software Engineer at XYZ Corp' in all_text
    assert 'Led team of 5 developers' in all_text
