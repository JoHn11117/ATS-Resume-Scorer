import pytest
from docx import Document
from docx.shared import Pt
from io import BytesIO
from backend.services.section_detector import SectionDetector

def test_detect_sections_by_heading_style():
    """Test detecting sections using Heading styles"""
    # Create test DOCX
    doc = Document()
    doc.add_heading('Experience', level=2)
    doc.add_paragraph('Software Engineer at ABC Corp')
    doc.add_heading('Education', level=2)
    doc.add_paragraph('BS Computer Science')

    # Save to bytes
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    # Detect sections
    detector = SectionDetector()
    sections = detector.detect(docx_bytes.read())

    # Assertions
    assert len(sections) == 2
    assert sections[0]['title'] == 'Experience'
    assert sections[0]['content'] == 'Software Engineer at ABC Corp'
    assert sections[0]['section_id'] == 'section_0'
    assert sections[0]['start_para_idx'] >= 0
    assert sections[0]['end_para_idx'] > sections[0]['start_para_idx']
    assert sections[1]['title'] == 'Education'

def test_detect_sections_by_bold_text():
    """Test detecting sections using bold text"""
    doc = Document()

    # Add bold heading
    p1 = doc.add_paragraph()
    run1 = p1.add_run('WORK EXPERIENCE')
    run1.bold = True
    run1.font.size = Pt(14)

    doc.add_paragraph('Senior Developer at XYZ')

    # Add another bold heading
    p2 = doc.add_paragraph()
    run2 = p2.add_run('SKILLS')
    run2.bold = True
    run2.font.size = Pt(14)

    doc.add_paragraph('Python, JavaScript, React')

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    detector = SectionDetector()
    sections = detector.detect(docx_bytes.read())

    assert len(sections) == 2
    assert sections[0]['title'] == 'WORK EXPERIENCE'
    assert 'Senior Developer' in sections[0]['content']
    assert sections[1]['title'] == 'SKILLS'

def test_detect_sections_by_all_caps():
    """Test detecting sections using ALL CAPS text"""
    doc = Document()
    doc.add_paragraph('PROFESSIONAL SUMMARY')
    doc.add_paragraph('Experienced software engineer with 5 years of experience')
    doc.add_paragraph('TECHNICAL SKILLS')
    doc.add_paragraph('Python, Java, React, Node.js')

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    detector = SectionDetector()
    sections = detector.detect(docx_bytes.read())

    assert len(sections) == 2
    assert sections[0]['title'] == 'PROFESSIONAL SUMMARY'
    assert sections[1]['title'] == 'TECHNICAL SKILLS'
    assert 'Experienced software engineer' in sections[0]['content']

def test_empty_document():
    """Test handling of empty document"""
    doc = Document()

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    detector = SectionDetector()
    sections = detector.detect(docx_bytes.read())

    assert len(sections) == 0

def test_only_headings_no_content():
    """Test document with only headings and no content"""
    doc = Document()
    doc.add_heading('Experience', level=2)
    doc.add_heading('Education', level=2)
    doc.add_heading('Skills', level=2)

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    detector = SectionDetector()
    sections = detector.detect(docx_bytes.read())

    # Should detect sections even with no content
    assert len(sections) == 3
    assert sections[0]['title'] == 'Experience'
    assert sections[0]['content'] == ''
    assert sections[1]['title'] == 'Education'
    assert sections[1]['content'] == ''

def test_invalid_docx_bytes():
    """Test handling of invalid DOCX bytes"""
    detector = SectionDetector()

    with pytest.raises(ValueError, match="Invalid DOCX format"):
        detector.detect(b'invalid docx data')

def test_empty_bytes():
    """Test handling of empty bytes"""
    detector = SectionDetector()

    with pytest.raises(ValueError, match="docx_bytes cannot be empty"):
        detector.detect(b'')

def test_detect_contact_section():
    """Test detecting Contact section"""
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("john@example.com")
    doc.add_paragraph("(555) 123-4567")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Software Engineer at TechCorp")

    detector = SectionDetector()
    sections = detector.detect_sections(doc)

    assert len(sections) > 0
    contact = next((s for s in sections if s['name'] == 'Contact'), None)
    assert contact is not None
    assert contact['start_para'] == 0
    assert contact['end_para'] >= 2
