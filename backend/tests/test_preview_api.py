import pytest
from fastapi.testclient import TestClient
from docx import Document
from io import BytesIO
import uuid
from backend.main import app
from backend.services.docx_template_manager import DocxTemplateManager

client = TestClient(app)

@pytest.fixture
def test_session():
    """Create test session with DOCX using default storage"""
    # Create test DOCX
    doc = Document()
    doc.add_heading('Test Resume', level=1)
    doc.add_paragraph('Test content')

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    # Use default template manager (same instance as API)
    manager = DocxTemplateManager()
    session_id = str(uuid.uuid4())
    manager.save_template(session_id, docx_bytes.read())

    yield session_id

    # Cleanup: remove test files
    import shutil
    original_path = manager.storage_dir / f"{session_id}_original.docx"
    working_path = manager.storage_dir / f"{session_id}_working.docx"
    if original_path.exists():
        original_path.unlink()
    if working_path.exists():
        working_path.unlink()

def test_get_preview_docx(test_session):
    """Test serving preview DOCX file"""
    response = client.get(f"/api/preview/{test_session}.docx")

    assert response.status_code == 200
    assert response.headers['content-type'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    assert len(response.content) > 0

def test_get_preview_invalid_session_id():
    """Test that invalid session IDs are rejected (security)"""
    # Test invalid UUID formats (path traversal attempts would be caught by routing)
    invalid_ids = [
        "invalid-uuid-format",
        "test@session",
        "12345",
        "not-a-uuid",
    ]

    for invalid_id in invalid_ids:
        response = client.get(f"/api/preview/{invalid_id}.docx")
        # Should return 400 for invalid UUID format
        assert response.status_code == 400, f"Expected 400 for {invalid_id}, got {response.status_code}"
        assert "Invalid session ID format" in response.json()["detail"]

def test_update_section_invalid_session_id():
    """Test that update rejects invalid session IDs (security)"""
    invalid_ids = [
        "../../../malicious",
        "invalid-uuid",
        "test@session",
    ]

    for invalid_id in invalid_ids:
        response = client.post("/api/preview/update", json={
            "session_id": invalid_id,
            "start_para_idx": 0,
            "end_para_idx": 0,
            "new_content": "test"
        })
        assert response.status_code == 400
        assert "Invalid session ID format" in response.json()["detail"]
