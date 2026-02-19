from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import uuid
from docx import Document
from bs4 import BeautifulSoup
from pathlib import Path
import logging

from backend.services.suggestion_generator import SuggestionGenerator
from backend.services.docx_template_manager import DocxTemplateManager

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/editor", tags=["editor"])

# In-memory session store (TODO: Replace with persistent storage)
SESSION_STORE: Dict[str, Dict] = {}

# Initialize template manager
template_manager = DocxTemplateManager()

class CreateSessionRequest(BaseModel):
    resume_id: str
    role: Optional[str] = "software_engineer"
    level: Optional[str] = "mid"

class SessionResponse(BaseModel):
    session_id: str
    working_docx_url: str
    sections: List[Dict[str, Any]]
    current_score: Dict[str, Any]
    suggestions: List[Dict[str, Any]]

class RescoreRequest(BaseModel):
    session_id: str

class RescoreResponse(BaseModel):
    score: Dict[str, Any]
    suggestions: List[Dict[str, Any]]


class UpdateSectionRequest(BaseModel):
    session_id: str
    section: str
    content: str  # HTML from TipTap
    start_para: int
    end_para: int


class UpdateSectionResponse(BaseModel):
    success: bool
    updated_url: str

@router.post("/session", response_model=SessionResponse)
async def create_editor_session(request: CreateSessionRequest):
    """Create new editing session from uploaded resume"""
    session_id = str(uuid.uuid4())

    # TODO: Get actual DOCX file for resume_id
    # For now, create sample document with some content to generate suggestions
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("john@example.com")  # Missing phone
    doc.add_paragraph("Experience")
    doc.add_paragraph("Responsible for managing team")  # Weak verb

    # Save the sample document to storage for testing
    storage_path = Path(__file__).parent.parent / "storage" / "templates"
    storage_path.mkdir(parents=True, exist_ok=True)

    # Save working copy
    working_path = storage_path / f"{session_id}_working.docx"
    doc.save(working_path)

    # Generate suggestions using SuggestionGenerator
    generator = SuggestionGenerator(role=request.role, level=request.level)

    # Create sample resume data for suggestion generation
    resume_data = {
        'contact': {
            'name': 'John Doe',
            'email': 'john@example.com'
            # Missing phone and linkedin
        },
        'experience': [
            {
                'description': 'Responsible for managing team',
                'para_idx': 3
            }
        ],
        'skills': [],
        'education': []
    }

    # Create sample sections
    sections = [
        {'name': 'Contact', 'start_para': 0, 'end_para': 1},
        {'name': 'Experience', 'start_para': 2, 'end_para': 3}
    ]

    suggestions = generator.generate_suggestions(resume_data, sections)

    response_data = SessionResponse(
        session_id=session_id,
        working_docx_url=f"/api/downloads/{session_id}_working.docx",
        sections=sections,
        current_score={"overallScore": 0},
        suggestions=suggestions
    )

    # Store session in memory
    SESSION_STORE[session_id] = {
        "session_id": session_id,
        "working_docx_url": response_data.working_docx_url,
        "sections": sections,
        "current_score": response_data.current_score,
        "suggestions": suggestions
    }

    return response_data

@router.get("/session/{session_id}", response_model=SessionResponse)
async def get_editor_session(session_id: str):
    """Get existing session state"""
    if session_id not in SESSION_STORE:
        raise HTTPException(status_code=404, detail="Session not found")

    session_data = SESSION_STORE[session_id]

    return SessionResponse(
        session_id=session_data["session_id"],
        working_docx_url=session_data["working_docx_url"],
        sections=session_data["sections"],
        current_score=session_data["current_score"],
        suggestions=session_data["suggestions"]
    )

@router.post("/rescore", response_model=RescoreResponse)
async def rescore_resume(request: RescoreRequest):
    """Re-score the current working DOCX and generate fresh suggestions"""
    from backend.services.scorer_ats import ATSScorer
    from backend.services.parser import ResumeData

    # Validate session exists
    if request.session_id not in SESSION_STORE:
        raise HTTPException(status_code=404, detail="Session not found")

    session_data = SESSION_STORE[request.session_id]

    # TODO: Load actual working DOCX from storage
    # For now, create a mock document
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("john@example.com")
    doc.add_paragraph("(555) 123-4567")  # Added phone
    doc.add_paragraph("Experience")
    doc.add_paragraph("Developed and managed team projects")  # Improved verb

    # Create mock resume data for scoring
    # In production, this should parse the actual DOCX
    mock_resume_data = ResumeData(
        fileName="test_resume.docx",
        contact={
            "name": "John Doe",
            "email": "john@example.com",
            "phone": "(555) 123-4567",
            "linkedin": "linkedin.com/in/johndoe",
            "location": "San Francisco, CA"
        },
        experience=[
            {
                "title": "Software Engineer",
                "company": "Tech Corp",
                "startDate": "2020-01",
                "endDate": "Present",
                "description": "Developed scalable web applications using Python and React, serving 50K+ users"
            }
        ],
        education=[
            {
                "degree": "BS Computer Science",
                "institution": "Tech University",
                "graduationDate": "2020"
            }
        ],
        skills=["Python", "JavaScript", "React", "AWS", "Docker"],
        certifications=[],
        metadata={
            "pageCount": 1,
            "wordCount": 350,
            "hasPhoto": False,
            "fileFormat": "docx"
        }
    )

    # Run ATS scorer with mock data
    scorer = ATSScorer()
    try:
        score_result = scorer.score(
            resume=mock_resume_data,
            role="software_engineer",
            level="mid",
            job_description=""
        )

        # Extract overall score and breakdown
        score = {
            "overallScore": score_result.get("score", 0),
            "breakdown": {}
        }

        # Add breakdown details
        breakdown = score_result.get("breakdown", {})
        for category, data in breakdown.items():
            score["breakdown"][category] = {
                "score": data.get("score", 0),
                "maxScore": data.get("maxScore", 0),
                "details": data.get("details", {})
            }
    except Exception as e:
        # Fallback to default score if scoring fails
        score = {
            "overallScore": 60,
            "breakdown": {
                "contact": {"score": 4, "maxScore": 5, "details": {}},
                "experience": {"score": 12, "maxScore": 20, "details": {}},
                "education": {"score": 12, "maxScore": 15, "details": {}},
                "skills": {"score": 10, "maxScore": 12, "details": {}}
            }
        }

    # Generate fresh suggestions
    generator = SuggestionGenerator(role="software_engineer", level="mid")
    sections = session_data.get("sections", [])

    # Create resume data dict for suggestion generation
    resume_data_dict = {
        "contact": mock_resume_data.contact,
        "experience": mock_resume_data.experience,
        "education": mock_resume_data.education,
        "skills": mock_resume_data.skills,
        "certifications": mock_resume_data.certifications
    }

    suggestions = generator.generate_suggestions(
        resume_data=resume_data_dict,
        sections=sections
    )

    # Update session with new score
    SESSION_STORE[request.session_id]["current_score"] = score
    SESSION_STORE[request.session_id]["suggestions"] = suggestions

    return RescoreResponse(
        score=score,
        suggestions=suggestions
    )


@router.post("/update-section", response_model=UpdateSectionResponse)
async def update_section(request: UpdateSectionRequest):
    """Update specific section in working DOCX"""

    # Check if session exists
    if not template_manager.working_exists(request.session_id):
        raise HTTPException(status_code=404, detail="Session not found")

    try:
        # Parse HTML content to extract text
        soup = BeautifulSoup(request.content, 'html.parser')
        text = soup.get_text(separator='\n').strip()

        # If empty after parsing, use original content
        if not text:
            text = request.content

        # Update section using template manager
        result = template_manager.update_section(
            session_id=request.session_id,
            start_para_idx=request.start_para,
            end_para_idx=request.end_para,
            new_content=text
        )

        if not result.get('success'):
            error_msg = result.get('error', 'Unknown error')
            raise HTTPException(status_code=400, detail=error_msg)

        # Return success with updated URL
        return UpdateSectionResponse(
            success=True,
            updated_url=f"/api/downloads/{request.session_id}_working.docx"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update section: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class ApplySuggestionRequest(BaseModel):
    session_id: str
    suggestion_id: str
    action: str
    value: Optional[str] = None


class ApplySuggestionResponse(BaseModel):
    success: bool
    updated_section: str
    content: str


# In-memory DOCX storage (TODO: Replace with file system or S3)
DOCX_STORE: Dict[str, Document] = {}


@router.post("/apply-suggestion", response_model=ApplySuggestionResponse)
async def apply_suggestion(request: ApplySuggestionRequest):
    """Apply a suggestion action to the working DOCX

    Supports 4 action types:
    1. add_phone - Add phone number to contact section
    2. replace_text - Replace weak text with stronger alternative
    3. add_section - Add new section with content
    4. show_location - Navigate to location (no-op for API)
    """
    import json

    # Get session
    if request.session_id not in SESSION_STORE:
        raise HTTPException(status_code=404, detail="Session not found")

    # Get or create DOCX document
    if request.session_id not in DOCX_STORE:
        # Create a sample document if not exists
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("john@example.com")
        doc.add_paragraph("Experience")
        doc.add_paragraph("Responsible for managing team")
        DOCX_STORE[request.session_id] = doc
    else:
        doc = DOCX_STORE[request.session_id]

    # Action: add_phone
    if request.action == "add_phone":
        # Add phone to first paragraph (contact section)
        if len(doc.paragraphs) > 1:
            para = doc.paragraphs[1]  # Usually contact is in para 1
            # Add phone as new line in contact section
            para.text = para.text + f"\nPhone: {request.value}"
        else:
            doc.add_paragraph(f"Phone: {request.value}")

        content = f"<p>Phone: {request.value}</p>"

        return ApplySuggestionResponse(
            success=True,
            updated_section="Contact",
            content=content
        )

    # Action: replace_text
    elif request.action == "replace_text":
        try:
            value_data = json.loads(request.value) if request.value else {}
            current_text = value_data.get("current_text", "")
            suggested_text = value_data.get("suggested_text", "")
            para_idx = value_data.get("para_idx", 0)

            # Find and replace text in specified paragraph
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                if current_text in para.text:
                    para.text = para.text.replace(current_text, suggested_text)

                    content = f"<p>{suggested_text}</p>"

                    return ApplySuggestionResponse(
                        success=True,
                        updated_section="Experience",
                        content=content
                    )

            raise HTTPException(status_code=400, detail="Text not found or paragraph index invalid")

        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid value format for replace_text")

    # Action: add_section
    elif request.action == "add_section":
        # Parse section content (format: "Section Title\nContent line 1\nContent line 2")
        lines = request.value.split("\n") if request.value else []
        if not lines:
            raise HTTPException(status_code=400, detail="No section content provided")

        section_title = lines[0]

        # Add section heading
        heading_para = doc.add_paragraph(section_title)
        heading_para.style = 'Heading 1'

        # Add section content
        for line in lines[1:]:
            if line.strip():
                doc.add_paragraph(line.strip())

        # Build HTML content
        content_lines = [f"<h2>{section_title}</h2>"]
        for line in lines[1:]:
            if line.strip():
                content_lines.append(f"<p>{line.strip()}</p>")

        content = "\n".join(content_lines)

        return ApplySuggestionResponse(
            success=True,
            updated_section=section_title,
            content=content
        )

    # Action: show_location (navigate only, no modification)
    elif request.action == "show_location":
        # This is a no-op for the API - frontend handles navigation
        return ApplySuggestionResponse(
            success=True,
            updated_section="",
            content=""
        )

    # Unknown action
    raise HTTPException(status_code=400, detail=f"Unknown action: {request.action}")
