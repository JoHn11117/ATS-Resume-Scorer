#!/usr/bin/env python3
"""
Comprehensive test runner for DocxTemplateManager.
Follows TDD approach: tests were written first, then implementation.
"""
import sys
import os

# Add backend directory to path
sys.path.insert(0, '/Users/sabuj.mondal/ats-resume-scorer/backend')

from docx import Document
from io import BytesIO
import tempfile
import shutil
from pathlib import Path
from services.docx_template_manager import DocxTemplateManager


def create_test_docx():
    """Create a test DOCX file"""
    doc = Document()
    doc.add_heading('Contact Information', level=2)
    doc.add_paragraph('John Doe')
    doc.add_paragraph('john@example.com')
    doc.add_heading('Experience', level=2)
    doc.add_paragraph('Software Engineer at ABC Corp')

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.read()


def test_save_template():
    """Test saving original template and creating working copy"""
    print("\n" + "="*60)
    print("TEST 1: test_save_template")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmp_path:
        print(f"Creating template manager with temp dir: {tmp_path}")
        template_manager = DocxTemplateManager(storage_dir=tmp_path)
        test_docx = create_test_docx()
        session_id = "test_session_123"

        print(f"Saving template for session: {session_id}")
        template_path = template_manager.save_template(session_id, test_docx)

        # Verify file exists
        print(f"Checking original template exists: {template_path}")
        assert os.path.exists(template_path), f"Template path does not exist: {template_path}"
        assert session_id in template_path, f"Session ID not in path: {template_path}"
        assert template_path.endswith('_original.docx'), f"Path doesn't end with _original.docx: {template_path}"

        # Verify working copy created
        working_path = template_path.replace('_original.docx', '_working.docx')
        print(f"Checking working copy exists: {working_path}")
        assert os.path.exists(working_path), f"Working copy does not exist: {working_path}"

        print("✓ test_save_template PASSED")
        print("  - Original template saved successfully")
        print("  - Working copy created successfully")
        return True


def test_update_section_content():
    """Test updating specific section content while preserving formatting"""
    print("\n" + "="*60)
    print("TEST 2: test_update_section_content")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmp_path:
        print(f"Creating template manager with temp dir: {tmp_path}")
        template_manager = DocxTemplateManager(storage_dir=tmp_path)
        test_docx = create_test_docx()
        session_id = "test_session_456"

        print(f"Saving template for session: {session_id}")
        template_manager.save_template(session_id, test_docx)

        # Read initial content
        working_path = template_manager.get_working_path(session_id)
        doc_before = Document(working_path)
        print(f"\nInitial paragraphs ({len(doc_before.paragraphs)} total):")
        for idx, para in enumerate(doc_before.paragraphs):
            print(f"  [{idx}] {para.text}")

        # Update Experience section (paragraph index 4)
        new_content = "Senior Software Engineer at XYZ Corp\nLed team of 5 developers"
        print(f"\nUpdating paragraph 4 with new content:")
        print(f"  {new_content}")

        result = template_manager.update_section(
            session_id=session_id,
            start_para_idx=4,
            end_para_idx=4,
            new_content=new_content
        )

        print(f"\nUpdate result: {result}")
        assert result['success'] is True, f"Update failed: {result.get('error', 'Unknown error')}"
        assert 'preview_url' in result, "Preview URL not in result"

        # Verify content updated
        doc_after = Document(working_path)
        print(f"\nUpdated paragraphs ({len(doc_after.paragraphs)} total):")
        for idx, para in enumerate(doc_after.paragraphs):
            print(f"  [{idx}] {para.text}")

        # Check that new content is in document
        all_text = '\n'.join([p.text for p in doc_after.paragraphs])
        assert 'Senior Software Engineer at XYZ Corp' in all_text, \
            "First line of new content not found in document"
        assert 'Led team of 5 developers' in all_text, \
            "Second line of new content not found in document"

        print("\n✓ test_update_section_content PASSED")
        print("  - Section updated successfully")
        print("  - New content inserted correctly")
        print("  - Preview URL generated")
        return True


def test_working_exists():
    """Test working_exists method"""
    print("\n" + "="*60)
    print("TEST 3: test_working_exists")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmp_path:
        template_manager = DocxTemplateManager(storage_dir=tmp_path)
        session_id = "test_session_789"

        # Should not exist initially
        assert not template_manager.working_exists(session_id), \
            "Working copy should not exist before saving"
        print("✓ Working copy correctly reported as non-existent")

        # Create template
        test_docx = create_test_docx()
        template_manager.save_template(session_id, test_docx)

        # Should exist now
        assert template_manager.working_exists(session_id), \
            "Working copy should exist after saving"
        print("✓ Working copy correctly reported as existing")

        print("\n✓ test_working_exists PASSED")
        return True


def test_get_working_path():
    """Test get_working_path method"""
    print("\n" + "="*60)
    print("TEST 4: test_get_working_path")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmp_path:
        template_manager = DocxTemplateManager(storage_dir=tmp_path)
        session_id = "test_session_path"

        working_path = template_manager.get_working_path(session_id)
        print(f"Working path: {working_path}")

        assert isinstance(working_path, Path), "Should return Path object"
        assert session_id in str(working_path), "Session ID should be in path"
        assert str(working_path).endswith('_working.docx'), "Should end with _working.docx"

        print("✓ test_get_working_path PASSED")
        return True


def run_all_tests():
    """Run all tests and report results"""
    print("\n" + "="*60)
    print("DOCX TEMPLATE MANAGER TEST SUITE")
    print("="*60)
    print("Testing: backend/services/docx_template_manager.py")
    print("Following TDD approach: tests written first, then implementation")

    tests = [
        ("test_save_template", test_save_template),
        ("test_update_section_content", test_update_section_content),
        ("test_working_exists", test_working_exists),
        ("test_get_working_path", test_get_working_path),
    ]

    passed = 0
    failed = 0
    errors = []

    for test_name, test_func in tests:
        try:
            test_func()
            passed += 1
        except AssertionError as e:
            failed += 1
            errors.append((test_name, str(e)))
            print(f"\n✗ {test_name} FAILED: {e}")
        except Exception as e:
            failed += 1
            errors.append((test_name, f"ERROR: {e}"))
            print(f"\n✗ {test_name} ERROR: {e}")
            import traceback
            traceback.print_exc()

    # Print summary
    print("\n" + "="*60)
    print("TEST SUMMARY")
    print("="*60)
    print(f"Total tests: {len(tests)}")
    print(f"Passed: {passed}")
    print(f"Failed: {failed}")

    if errors:
        print("\nFailed tests:")
        for test_name, error in errors:
            print(f"  - {test_name}: {error}")

    if failed == 0:
        print("\n✓ All tests PASSED!")
        print("\nImplementation complete:")
        print("  - DocxTemplateManager.save_template() ✓")
        print("  - DocxTemplateManager.update_section() ✓")
        print("  - DocxTemplateManager.get_working_path() ✓")
        print("  - DocxTemplateManager.working_exists() ✓")
        return 0
    else:
        print(f"\n✗ {failed} test(s) failed")
        return 1


if __name__ == "__main__":
    sys.exit(run_all_tests())
